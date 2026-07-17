"""Document processing service using Azure Document Intelligence.

This module provides document parsing and processing capabilities using Azure Document Intelligence.  # noqa: E501
It handles various document formats including PDF, DOCX, Markdown, and plain text.

Example:
    ```python
    # Initialize service
    doc_service = DocumentService()

    # Parse a PDF file
    with open('document.pdf', 'rb') as f:
        content, metadata = doc_service.parse_file(f, 'application/pdf')
    ```
"""

# Removed: import logging
import time

# Removed: from functools import wraps
from typing import (  # noqa: E501, F401
    Any,
    BinaryIO,
    Dict,
    List,
    Optional,
    Tuple,
    TypeVar,
)

import docx
import markdown
from azure.core.exceptions import AzureError, ResourceExistsError
from bs4 import BeautifulSoup

# Removed ServiceLoggingMixin, AzureClientMixin imports
from tenacity import (
    retry,
    retry_if_exception_type,
    stop_after_attempt,
    wait_exponential,
)

from konveyor.core.azure_utils.clients import (  # Import client manager  # noqa: F401
    AzureClientManager,
)

# Removed AzureKeyCredential, DocumentIntelligenceClient imports
from konveyor.core.azure_utils.service import AzureService  # Import base service

# Removed module-level logger

# Removed duplicate mixin import


class DocumentService(AzureService):  # Inherit from AzureService
    """Service for handling document operations.

    This service provides methods for:
    - Parsing different document types
    - Extracting text and metadata
    - Managing document storage

    It uses Azure Document Intelligence for document processing.
    """

    # Removed redundant log_warning method (provided by AzureService base class)
    """Service for processing documents using Azure Document Intelligence.
      # noqa: W293
    This service provides methods to parse and process various document formats using
    Azure Document Intelligence. It supports PDF, DOCX, Markdown, and plain text files.
      # noqa: W293
    Attributes:
        doc_intelligence_client (DocumentIntelligenceClient): Azure Document Intelligence client  # noqa: E501
    Required Environment Variables:
        AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT: Document Intelligence endpoint
        AZURE_DOCUMENT_INTELLIGENCE_API_KEY: Document Intelligence API key
    """

    def __init__(self):
        """Initialize document service using AzureService base."""
        # Call parent constructor - handles config, client manager, validation
        super().__init__("DOCUMENT_INTELLIGENCE")
        self.log_init("DocumentService")  # Keep specific component log

        # Get clients from the manager inherited from AzureService
        # Ensure 'DOCUMENT_INTELLIGENCE' is configured in AzureConfig validation map
        self.doc_intelligence_client = (
            self.client_manager.get_document_intelligence_client()
        )
        # Blob client is retrieved on demand in storage methods using self.client_manager  # noqa: E501

        self.log_success("DocumentService initialized using AzureService base")

    # Removed initialize_document_intelligence_client method

    def parse_file(
        self, file_obj: BinaryIO, content_type: str
    ) -> tuple[str, dict[str, Any]]:
        """Parse a document file and extract its content and metadata.

        Args:
            file_obj (BinaryIO): File-like object containing the document
            content_type (str): MIME type of the document. Supported types:
                - application/pdf
                - application/vnd.openxmlformats-officedocument.wordprocessingml.document  # noqa: E501
                - text/markdown
                - text/plain

        Returns:
            Tuple[str, Dict[str, Any]]: A tuple containing:
                - Extracted text content
                - Document metadata

        Raises:
            ValueError: If content type is not supported
            Exception: If parsing fails
        """
        try:
            if content_type == "application/pdf":
                return self._parse_pdf(file_obj)
            elif (
                content_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"  # noqa: E501
            ):
                return self._parse_docx(file_obj)
            elif content_type == "text/markdown":
                return self._parse_markdown(file_obj)
            elif content_type == "text/plain":
                return self._parse_text(file_obj)
            else:
                error_msg = f"Unsupported content type: {content_type}"
                self.log_error(error_msg)
                raise ValueError(error_msg)
        except Exception as e:
            self.log_error(f"Failed to parse file with content type {content_type}", e)
            raise

    @retry(
        retry=retry_if_exception_type(AzureError),
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
    )
    def _parse_pdf(self, file_obj: BinaryIO) -> tuple[str, dict[str, Any]]:
        """Parse PDF using Azure Document Intelligence.

        Args:
            file_obj (BinaryIO): PDF file object

        Returns:
            Tuple[str, Dict[str, Any]]: Extracted text and metadata

        Raises:
            Exception: If PDF parsing fails
        """
        try:
            # Get model name from config or use default
            model_name = self.config.get_setting(
                "AZURE_DOCUMENT_INTELLIGENCE_MODEL", default="prebuilt-layout"
            )
            self.log_debug(
                f"Using Document Intelligence model: {model_name}"
            )  # Use debug level

            poller = self.doc_intelligence_client.begin_analyze_document(
                model_name,
                body=file_obj,
                content_type="application/pdf",  # Explicitly set content type
            )
            result = poller.result()

            content = []
            for page in result.pages:
                if page.lines:
                    content.extend(line.content for line in page.lines)

            metadata = {
                "page_count": len(result.pages),
                "language": result.languages[0] if result.languages else None,
                "document_type": "pdf",
            }

            self.log_success(
                f"Successfully parsed PDF with {metadata['page_count']} pages"
            )
            return "\n".join(content), metadata

        except Exception as e:
            self.log_error("Failed to parse PDF", e)
            raise

    @retry(
        retry=retry_if_exception_type(AzureError),
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
    )
    def _parse_docx(self, file_obj: BinaryIO) -> tuple[str, dict[str, Any]]:
        """Parse DOCX file.

        Args:
            file_obj (BinaryIO): DOCX file object

        Returns:
            Tuple[str, Dict[str, Any]]: Extracted text and metadata

        Raises:
            Exception: If DOCX parsing fails
        """
        try:
            doc = docx.Document(file_obj)
            content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)

            metadata = {"paragraph_count": len(doc.paragraphs), "document_type": "docx"}

            self.log_success(
                f"Successfully parsed DOCX with {metadata['paragraph_count']} paragraphs"  # noqa: E501
            )
            return "\n".join(content), metadata

        except Exception as e:
            self.log_error("Failed to parse DOCX", e)
            raise

    @retry(
        retry=retry_if_exception_type(AzureError),
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
    )
    def _parse_markdown(self, file_obj: BinaryIO) -> tuple[str, dict[str, Any]]:
        """Parse Markdown file.

        Args:
            file_obj (BinaryIO): Markdown file object

        Returns:
            Tuple[str, Dict[str, Any]]: Raw markdown content and metadata

        Raises:
            Exception: If Markdown parsing fails
        """
        try:
            content = file_obj.read().decode("utf-8")
            html = markdown.markdown(content)
            soup = BeautifulSoup(html, "html.parser")

            metadata = {
                "document_type": "markdown",
                "has_headers": bool(
                    soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"])
                ),
            }

            self.log_success("Successfully parsed Markdown")
            return content, metadata

        except Exception as e:
            self.log_error("Failed to parse Markdown", e)
            raise

    @retry(
        retry=retry_if_exception_type(AzureError),
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
    )
    def _parse_text(self, file_obj: BinaryIO) -> tuple[str, dict[str, Any]]:
        """Parse plain text file.

        Args:
            file_obj (BinaryIO): Text file object

        Returns:
            Tuple[str, Dict[str, Any]]: Text content and metadata

        Raises:
            Exception: If text parsing fails
        """
        try:
            content = file_obj.read().decode("utf-8")
            lines = content.splitlines()

            metadata = {"document_type": "text", "line_count": len(lines)}

            self.log_success(
                f"Successfully parsed text file with {metadata['line_count']} lines"
            )
            return content, metadata

        except Exception as e:
            self.log_error("Failed to parse text file", e)
            raise

    def store_chunk_content(self, chunk, content: str) -> None:
        """Store document chunk content in Azure Blob Storage.

        Args:
            chunk: DocumentChunk instance to store content for
            content (str): Content to store

        Raises:
            ValueError: If chunk has no blob_path
            Exception: If storage fails
        """
        if not chunk.blob_path:
            raise ValueError("Chunk has no blob_path")

        try:
            blob_service_client = self.client_manager.get_blob_client()
            # Use test container name from env var if set, else default
            container_name = self.config.get_setting(
                "AZURE_STORAGE_CONTAINER_NAME", default="document-chunks"
            )
            blob_name = chunk.blob_path

            # Get container client
            container_client = blob_service_client.get_container_client(container_name)

            # Create container if it doesn't exist with specific retry for ContainerBeingDeleted  # noqa: E501
            max_retries = 5
            retry_delay = 2
            container_exists_or_created = False

            for attempt in range(max_retries):
                try:
                    if container_client.exists():
                        self.log_debug(
                            f"Container {container_name} already exists."
                        )  # Use debug
                        container_exists_or_created = True
                        break
                    else:
                        self.log_debug(
                            f"Attempting to create container {container_name}..."
                        )  # Use debug
                        container_client.create_container()
                        self.log_success(
                            f"Container {container_name} created successfully."
                        )
                        container_exists_or_created = True
                        break  # Exit loop on successful creation
                except ResourceExistsError as e:
                    # Check if the error is specifically "ContainerBeingDeleted"
                    if "ContainerBeingDeleted" in str(e):
                        if attempt == max_retries - 1:
                            self.log_error(
                                f"Failed to create container {container_name} after {max_retries} attempts due to ContainerBeingDeleted.",  # noqa: E501
                                e,
                            )
                            raise  # Re-raise the final error if all retries fail
                        self.log_warning(
                            f"Attempt {attempt + 1} failed: Container {container_name} is being deleted. Retrying in {retry_delay}s..."  # noqa: E501
                        )
                        # import time # Moved import to top level
                        time.sleep(retry_delay)
                        retry_delay *= 2  # Exponential backoff
                    else:
                        # If it's another ResourceExistsError (like ContainerAlreadyExists, which is fine if caught by exists()),  # noqa: E501
                        # or some other unexpected ResourceExistsError, log and re-raise immediately.  # noqa: E501
                        if "ContainerAlreadyExists" in str(e):
                            # This case should ideally be caught by container_client.exists(), but handle defensively  # noqa: E501
                            self.log_debug(
                                f"Container {container_name} already exists (caught in exception). Proceeding."  # noqa: E501
                            )  # Use debug
                            container_exists_or_created = True
                            break  # Container exists, proceed
                        else:
                            self.log_error(
                                f"Unexpected ResourceExistsError during container creation check/attempt",  # noqa: E501, F541
                                e,
                            )
                            raise  # Re-raise unexpected ResourceExistsError
                except Exception as e:
                    # Catch any other unexpected exceptions during container creation/check  # noqa: E501
                    self.log_error(
                        f"Unexpected error checking or creating container {container_name} on attempt {attempt + 1}",  # noqa: E501
                        e,
                    )
                    if attempt == max_retries - 1:
                        raise  # Re-raise after final attempt
                    # import time # Moved import to top level
                    time.sleep(retry_delay)
                    retry_delay *= 2

            if not container_exists_or_created:
                raise Exception(
                    f"Failed to ensure container {container_name} exists after {max_retries} attempts."  # noqa: E501
                )

            # Get blob client for this specific chunk within the container
            blob_client = container_client.get_blob_client(blob_name)

            # Upload content with retry (existing retry logic seems okay for upload)
            max_retries_upload = 3
            retry_delay_upload = 1
            upload_successful = False
            for attempt in range(max_retries_upload):
                try:
                    blob_client.upload_blob(content.encode("utf-8"), overwrite=True)
                    self.log_success(
                        f"Successfully stored content for chunk {chunk.id} in {container_name}/{blob_name}"  # noqa: E501
                    )
                    upload_successful = True
                    break
                except Exception as e:
                    if attempt == max_retries_upload - 1:
                        self.log_error(
                            f"Failed to upload blob {blob_name} after {max_retries_upload} attempts.",  # noqa: E501
                            e,
                        )
                        raise
                    self.log_warning(
                        f"Attempt {attempt + 1} failed to upload blob {blob_name}: {str(e)}. Retrying in {retry_delay_upload}s..."  # noqa: E501
                    )
                    # import time # Moved import to top level
                    time.sleep(retry_delay_upload)
                    retry_delay_upload *= 2

            if not upload_successful:
                raise Exception(
                    f"Failed to upload blob {blob_name} after {max_retries_upload} attempts."  # noqa: E501
                )

        except Exception as e:
            self.log_error(f"Failed to store content for chunk {chunk.id}: {str(e)}", e)
            raise

    def get_chunk_content(self, chunk) -> str:
        """Retrieve document chunk content from Azure Blob Storage.

        Args:
            chunk: DocumentChunk instance to retrieve content for

        Returns:
            str: Retrieved content

        Raises:
            ValueError: If chunk has no blob_path
            Exception: If retrieval fails
        """
        try:
            if not chunk.blob_path:
                raise ValueError("Chunk has no blob_path")

            # Get blob client
            blob_service_client = self.client_manager.get_blob_client()
            # Use test container name from env var if set, else default
            container_name = self.config.get_setting(
                "AZURE_STORAGE_CONTAINER_NAME", default="document-chunks"
            )
            blob_name = chunk.blob_path

            # Get container client
            container_client = blob_service_client.get_container_client(container_name)

            # Get blob client for this chunk
            blob_client = container_client.get_blob_client(blob_name)

            # Download content with retry
            max_retries = 3
            retry_delay = 1  # Initial delay in seconds

            for attempt in range(max_retries):
                try:
                    content = blob_client.download_blob().readall().decode("utf-8")
                    self.log_success(
                        f"Successfully retrieved content for chunk {chunk.id}"
                    )
                    return content
                except Exception as e:
                    if attempt == max_retries - 1:
                        raise
                    self.log_warning(
                        f"Attempt {attempt + 1} failed to download blob: {str(e)}. Retrying..."  # noqa: E501
                    )
                    # import time # Moved import to top level
                    time.sleep(retry_delay)
                    retry_delay *= 2

        except Exception as e:
            self.log_error(f"Failed to retrieve content for chunk {chunk.id}", e)
            raise
